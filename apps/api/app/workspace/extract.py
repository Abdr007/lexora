"""Turn whatever the user brought into pages of text.

One shape out — :class:`ExtractedDocument` — regardless of what went in, so everything
downstream (chunking, retrieval, reranking, citation verification) is the same code that
serves the law corpus. The formats differ only here.

    PDF with a text layer   PyMuPDF, the same reader the corpus uses
    PDF without one         rendered per page and sent to OCR
    DOCX                    python-docx, headings kept as structure
    TXT / MD                decoded, headings inferred from markdown
    HTML / URL              fetched, stripped, with SSRF defences below
    Images                  OCR

**OCR is deliberately two-tier.** Tesseract always works and costs nothing; Claude's
vision model is markedly better on a photographed page and costs a call. That is the same
split the rest of the project already makes between `offline-extractive` and
`Claude · grounded`, and it is labelled the same way, so a demo without a key degrades
visibly rather than silently.

Nothing here trusts its input. Every reader is wrapped so a malformed file returns a
message a user can act on instead of a 500, and the URL fetcher treats the URL as hostile
by default — see :func:`_assert_public_url`.
"""

from __future__ import annotations

import io
import ipaddress
import logging
import re
import socket
import unicodedata
from dataclasses import dataclass, field
from html.parser import HTMLParser
from typing import Any, Final
from urllib.parse import urlparse

from app.core.settings import Settings, get_settings

logger = logging.getLogger(__name__)

PDF_MAGIC: Final = b"%PDF"
DOCX_MAGIC: Final = b"PK\x03\x04"

IMAGE_SUFFIXES: Final = frozenset({".png", ".jpg", ".jpeg", ".webp", ".tif", ".tiff", ".bmp"})
TEXT_SUFFIXES: Final = frozenset({".txt", ".md", ".markdown", ".rst", ".csv", ".log"})
HTML_SUFFIXES: Final = frozenset({".html", ".htm"})

# A scanned page usually yields a handful of stray glyphs rather than nothing at all, so
# "has a text layer" cannot be `text != ""`. Below this many characters per page the PDF
# is treated as an image of a document.
TEXT_LAYER_CHARS_PER_PAGE: Final = 40


class ExtractionError(Exception):
    """Input that could not be read. The message is shown to the user verbatim."""


@dataclass(frozen=True, slots=True)
class ExtractedPage:
    page_no: int
    text: str


@dataclass(frozen=True, slots=True)
class ExtractedDocument:
    title: str
    source: str
    kind: str
    pages: tuple[ExtractedPage, ...]
    ocr_engine: str | None = None
    diagnostics: dict[str, Any] = field(default_factory=dict)

    @property
    def text(self) -> str:
        return "\n\n".join(page.text for page in self.pages if page.text.strip())


# ── text hygiene ─────────────────────────────────────────────────────────────


def normalise(text: str) -> str:
    """Fold the encoding damage that arrives with real documents.

    NFKC collapses the ligatures and full-width forms that PDF producers emit; the corpus
    parser fixes the same class of defect (AUDIT.md records a broken ligature that reached
    the BM25 index). Zero-width characters are dropped outright: they are invisible, they
    survive into a chunk, and they silently break exact-term matching.
    """
    cleaned = unicodedata.normalize("NFKC", text)
    cleaned = cleaned.replace("\u200b", "").replace("﻿", "").replace("\xad", "")
    cleaned = cleaned.replace("\r\n", "\n").replace("\r", "\n")
    cleaned = re.sub(r"[ \t]+", " ", cleaned)
    return re.sub(r"\n{3,}", "\n\n", cleaned).strip()


# ── HTML ─────────────────────────────────────────────────────────────────────


class _Readable(HTMLParser):
    """Strip a page to readable text.

    Written against the standard library rather than adding a parser dependency: the job
    is to drop script/style/nav and keep block structure, which does not need a full DOM.
    """

    _SKIP: Final = frozenset({"script", "style", "noscript", "svg", "head", "nav", "footer"})
    _BLOCK: Final = frozenset(
        {"p", "div", "br", "li", "tr", "section", "article", "h1", "h2", "h3", "h4", "h5", "h6"}
    )

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.parts: list[str] = []
        self.title: str = ""
        self._skip_depth = 0
        self._in_title = False

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        del attrs  # the readable-text pass needs tag structure, not attributes
        if tag in self._SKIP:
            self._skip_depth += 1
        elif tag == "title":
            self._in_title = True
        elif tag in self._BLOCK:
            self.parts.append("\n")

    def handle_endtag(self, tag: str) -> None:
        if tag in self._SKIP:
            self._skip_depth = max(0, self._skip_depth - 1)
        elif tag == "title":
            self._in_title = False
        elif tag in self._BLOCK:
            self.parts.append("\n")

    def handle_data(self, data: str) -> None:
        if self._in_title:
            self.title += data
        elif self._skip_depth == 0:
            self.parts.append(data)


def html_to_text(html: str) -> tuple[str, str]:
    """Return ``(title, text)``."""
    parser = _Readable()
    parser.feed(html)
    return parser.title.strip(), normalise("".join(parser.parts))


# ── URL fetching ─────────────────────────────────────────────────────────────


def _assert_public_url(url: str) -> None:
    """Refuse anything that could reach the host's own network.

    The server fetches a URL chosen by a stranger, which is a server-side request forgery
    primitive unless it is constrained. Every address the hostname resolves to is checked,
    not merely the first: a name that returns one public and one loopback address would
    otherwise pass this and then connect to the loopback one.

    This is defence in depth and not a complete answer — it cannot stop a DNS record that
    changes between this check and the connection. The size cap and the redirect ceiling
    in :func:`fetch_url` bound what a successful attempt could achieve.
    """
    parsed = urlparse(url)
    if parsed.scheme not in {"http", "https"}:
        raise ExtractionError(f"only http and https URLs are accepted, not {parsed.scheme!r}")
    host = parsed.hostname
    if not host:
        raise ExtractionError("that URL has no host")

    try:
        resolved = socket.getaddrinfo(
            host, parsed.port or (443 if parsed.scheme == "https" else 80)
        )
    except socket.gaierror as exc:
        raise ExtractionError(f"could not resolve {host}") from exc

    for info in resolved:
        address = ipaddress.ip_address(str(info[4][0]))
        if (
            address.is_private
            or address.is_loopback
            or address.is_link_local  # 169.254.169.254 is the cloud metadata endpoint
            or address.is_reserved
            or address.is_multicast
            or address.is_unspecified
        ):
            raise ExtractionError(
                f"{host} resolves to {address}, which is not a public address. "
                "Only public URLs can be fetched."
            )


def fetch_url(url: str, settings: Settings | None = None) -> ExtractedDocument:
    import httpx

    cfg = settings or get_settings()
    limit = cfg.workspace_max_bytes
    _assert_public_url(url)

    try:
        with httpx.Client(
            follow_redirects=False,
            timeout=cfg.workspace_fetch_timeout_s,
            headers={"User-Agent": "Lexora/1.0 (+document ingest)"},
        ) as client:
            response = client.get(url)
            # Redirects are followed by hand so each hop is re-validated. `follow_redirects`
            # would check only the URL the user supplied and then happily chase a 302 into
            # the private network the check above exists to keep it out of.
            for _ in range(cfg.workspace_max_redirects):
                if response.status_code not in {301, 302, 303, 307, 308}:
                    break
                location = response.headers.get("location", "")
                if not location:
                    break
                nxt = str(httpx.URL(url).join(location))
                _assert_public_url(nxt)
                url, response = nxt, client.get(nxt)
            response.raise_for_status()
    except httpx.HTTPStatusError as exc:
        raise ExtractionError(f"{url} returned HTTP {exc.response.status_code}") from exc
    except httpx.HTTPError as exc:
        raise ExtractionError(f"could not fetch {url}: {exc}") from exc

    body = response.content[: limit + 1]
    if len(body) > limit:
        raise ExtractionError(f"{url} is larger than the {limit // 1_000_000} MB limit")

    content_type = response.headers.get("content-type", "").split(";")[0].strip()
    if content_type == "application/pdf" or body[:4] == PDF_MAGIC:
        document = extract_bytes(body, filename=url.rsplit("/", 1)[-1] or "page.pdf", settings=cfg)
        return ExtractedDocument(
            title=document.title,
            source=url,
            kind="pdf",
            pages=document.pages,
            ocr_engine=document.ocr_engine,
            diagnostics=document.diagnostics,
        )

    charset = response.encoding or "utf-8"
    decoded = body.decode(charset, errors="replace")
    if "html" in content_type or decoded.lstrip()[:1] == "<":
        title, text = html_to_text(decoded)
    else:
        title, text = "", normalise(decoded)

    if not text.strip():
        raise ExtractionError(f"no readable text found at {url}")
    return ExtractedDocument(
        title=title or url,
        source=url,
        kind="html",
        pages=(ExtractedPage(1, text),),
        diagnostics={"content_type": content_type, "bytes": len(body)},
    )


# ── OCR ──────────────────────────────────────────────────────────────────────


def _ocr_tesseract(image_bytes: bytes) -> str | None:
    """Return text, or ``None`` when Tesseract is not installed in this image."""
    try:
        import pytesseract
        from PIL import Image
    except ImportError:
        return None
    try:
        with Image.open(io.BytesIO(image_bytes)) as image:
            return str(pytesseract.image_to_string(image))
    except Exception as exc:
        logger.warning("tesseract failed: %s", exc)
        return None


def _ocr_claude(image_bytes: bytes, media_type: str, settings: Settings) -> str | None:
    """Transcribe with Claude's vision model. ``None`` when no key is configured."""
    if not settings.use_anthropic:
        return None
    import base64

    from app.core.claude import get_client

    try:
        client = get_client(settings)
        if client is None:
            return None
        message = client.messages.create(
            model=settings.answer_model,
            max_tokens=4096,
            temperature=0,
            messages=[
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "image",
                            "source": {
                                "type": "base64",
                                "media_type": media_type,
                                "data": base64.b64encode(image_bytes).decode(),
                            },
                        },
                        {
                            "type": "text",
                            # Transcribe, do not summarise: this text becomes a citation,
                            # and a citation that paraphrases its source is worse than no
                            # citation at all.
                            "text": (
                                "Transcribe every word of this document image exactly as "
                                "written, preserving headings, numbering and paragraph "
                                "breaks. Do not summarise, explain, or add commentary. "
                                "If the image contains no legible text, reply with "
                                "exactly: NO_TEXT"
                            ),
                        },
                    ],
                }
            ],
        )
        parts = [block.text for block in message.content if getattr(block, "type", "") == "text"]
        text = "\n".join(parts).strip()
    except Exception as exc:
        logger.warning("claude vision OCR failed: %s", exc)
        return None
    # NO_TEXT is the model's own signal that the image is illegible. Returning None lets
    # the caller try Tesseract, which occasionally reads what the model declined to.
    return None if text == "NO_TEXT" else text


def ocr(image_bytes: bytes, media_type: str, settings: Settings) -> tuple[str, str]:
    """Return ``(text, engine)``. Claude when a key is present, else Tesseract."""
    text = _ocr_claude(image_bytes, media_type, settings)
    if text:
        return normalise(text), "claude-vision"
    text = _ocr_tesseract(image_bytes)
    if text:
        return normalise(text), "tesseract"
    raise ExtractionError(
        "no text could be read from that image. Claude vision needs "
        "LEXORA_ANTHROPIC_API_KEY, and Tesseract is not installed in this container."
    )


# ── PDF ──────────────────────────────────────────────────────────────────────


def _extract_pdf(
    data: bytes, settings: Settings
) -> tuple[list[ExtractedPage], str | None, dict[str, Any]]:
    import fitz

    try:
        document = fitz.open(stream=data, filetype="pdf")
    except Exception as exc:
        raise ExtractionError(f"that PDF could not be opened: {exc}") from exc

    with document:
        if document.needs_pass:
            raise ExtractionError("that PDF is password protected")
        page_count = min(document.page_count, settings.workspace_max_pages)
        raw = [(index + 1, document[index].get_text("text")) for index in range(page_count)]
        total_chars = sum(len(text) for _, text in raw)

        # A scanned PDF has pages but effectively no text layer. Render and OCR instead of
        # indexing a handful of stray glyphs and calling the document empty.
        if page_count and total_chars < TEXT_LAYER_CHARS_PER_PAGE * page_count:
            pages: list[ExtractedPage] = []
            engine: str | None = None
            for index in range(page_count):
                pixmap = document[index].get_pixmap(dpi=200)
                text, engine = ocr(pixmap.tobytes("png"), "image/png", settings)
                pages.append(ExtractedPage(index + 1, text))
            return (
                pages,
                engine,
                {
                    "pages": page_count,
                    "text_layer": False,
                    "truncated": document.page_count > page_count,
                },
            )

        return (
            [ExtractedPage(no, normalise(text)) for no, text in raw],
            None,
            {
                "pages": page_count,
                "text_layer": True,
                "truncated": document.page_count > page_count,
            },
        )


# ── DOCX ─────────────────────────────────────────────────────────────────────


def _extract_docx(data: bytes) -> list[ExtractedPage]:
    try:
        import docx
    except ImportError as exc:  # pragma: no cover - declared dependency
        raise ExtractionError("DOCX support is not installed in this container") from exc
    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:
        raise ExtractionError(f"that DOCX could not be opened: {exc}") from exc

    lines: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        # Word headings become markdown so the chunker's heading detection sees the
        # structure the author intended rather than an undifferentiated wall of text.
        style = (paragraph.style.name or "").lower() if paragraph.style else ""
        if style.startswith("heading"):
            depth = "".join(ch for ch in style if ch.isdigit()) or "1"
            lines.append(f"{'#' * min(int(depth), 6)} {text}")
        else:
            lines.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                lines.append(" | ".join(cells))

    text = normalise("\n\n".join(lines))
    if not text:
        raise ExtractionError("that DOCX contains no readable text")
    # Word has no fixed pages without rendering; one page keeps citations honest rather
    # than inventing page numbers that do not correspond to anything.
    return [ExtractedPage(1, text)]


# ── entry point ──────────────────────────────────────────────────────────────


def extract_bytes(
    data: bytes,
    filename: str,
    content_type: str | None = None,
    settings: Settings | None = None,
) -> ExtractedDocument:
    """Read an uploaded file. Dispatch is by magic bytes first, extension second."""
    cfg = settings or get_settings()
    if not data:
        raise ExtractionError("that file is empty")
    if len(data) > cfg.workspace_max_bytes:
        raise ExtractionError(
            f"that file is larger than the {cfg.workspace_max_bytes // 1_000_000} MB limit"
        )

    name = filename.strip() or "document"
    suffix = ("." + name.rsplit(".", 1)[-1].lower()) if "." in name else ""
    title = name.rsplit("/", 1)[-1]

    # Content sniffing before the extension: a `.txt` that is really a PDF should be read
    # as a PDF, and an extension is a claim by whoever uploaded the file.
    if data[:4] == PDF_MAGIC:
        pages, engine, diagnostics = _extract_pdf(data, cfg)
        document = ExtractedDocument(title, name, "pdf", tuple(pages), engine, diagnostics)
    elif data[:4] == DOCX_MAGIC and suffix == ".docx":
        document = ExtractedDocument(title, name, "docx", tuple(_extract_docx(data)))
    elif suffix in IMAGE_SUFFIXES or (content_type or "").startswith("image/"):
        media = content_type if content_type and content_type.startswith("image/") else "image/png"
        text, engine = ocr(data, media, cfg)
        document = ExtractedDocument(title, name, "image", (ExtractedPage(1, text),), engine)
    elif suffix in HTML_SUFFIXES:
        html_title, text = html_to_text(data.decode("utf-8", errors="replace"))
        document = ExtractedDocument(html_title or title, name, "html", (ExtractedPage(1, text),))
    elif suffix in TEXT_SUFFIXES or not suffix:
        text = normalise(data.decode("utf-8", errors="replace"))
        document = ExtractedDocument(title, name, "text", (ExtractedPage(1, text),))
    else:
        raise ExtractionError(
            f"{suffix or 'that file type'} is not supported. Upload a PDF, DOCX, image, "
            "HTML or plain-text file, or paste a link."
        )

    if not document.text.strip():
        raise ExtractionError(
            f"no readable text was found in {title}. If it is a scan, OCR produced nothing."
        )
    return document
